# -*- coding: utf-8 -*-
# pylint: disable=protected-access,redefined-outer-name
"""Unit tests for the workspace convert-office endpoint and page-break logic.

Tests cover:
- ``_get_docx_page_info``: parsing page dimensions from DOCX XML
- ``_has_page_break``: detecting explicit page breaks in mammoth's document model
- ``_estimate_element_lines``: estimating visual lines for paragraphs/tables
- ``_build_docx_transform``: building a transform that inserts page-break markers
- ``_convert_docx_to_html``: full DOCX → HTML conversion with page breaks
- Unsupported file-type handling
"""
from __future__ import annotations

import io
import zipfile
from pathlib import Path
from unittest.mock import patch

import pytest

from qwenpaw.app.routers.workspace import (
    _get_docx_page_info,
    _has_page_break,
    _estimate_element_lines,
    _build_docx_transform,
    _convert_docx_to_html,
    _PAGE_BREAK_MARKER,
)


# ─── Fixtures ──────────────────────────────────────────────────────────────


def _make_minimal_docx(tmp_path: Path, filename: str = "test.docx") -> Path:
    """Create a minimal .docx file using python-docx."""
    from docx import Document

    doc = Document()
    doc.add_heading("Test Document", level=0)
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("Hello World")
    doc.save(str(tmp_path / filename))
    return tmp_path / filename


def _make_long_docx(tmp_path: Path, filename: str = "long.docx") -> Path:
    """Create a .docx file with enough content to trigger estimated page breaks."""
    from docx import Document

    doc = Document()
    doc.add_heading("Long Document", level=0)
    # Add ~60 paragraphs to exceed one page (~45 lines)
    for i in range(60):
        doc.add_paragraph(f"Paragraph number {i} " * 10)
    doc.save(str(tmp_path / filename))
    return tmp_path / filename


def _make_docx_with_explicit_page_break(
    tmp_path: Path, filename: str = "explicit.docx"
) -> Path:
    """Create a .docx file with an explicit page break."""
    from docx import Document
    from docx.enum.text import WD_BREAK

    doc = Document()
    doc.add_paragraph("Page 1 content")
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Page 2 content")
    doc.save(str(tmp_path / filename))
    return tmp_path / filename


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    return _make_minimal_docx(tmp_path)


@pytest.fixture
def long_docx_file(tmp_path: Path) -> Path:
    return _make_long_docx(tmp_path)


@pytest.fixture
def explicit_break_docx(tmp_path: Path) -> Path:
    return _make_docx_with_explicit_page_break(tmp_path)


# ─── _get_docx_page_info ───────────────────────────────────────────────────


class TestGetDocxPageInfo:
    """Tests for ``_get_docx_page_info``."""

    def test_returns_defaults_for_nonexistent_file(self) -> None:
        info = _get_docx_page_info("/nonexistent/file.docx")
        assert "lines_per_page" in info
        assert "chars_per_line" in info
        assert info["lines_per_page"] > 0
        assert info["chars_per_line"] > 0

    def test_returns_defaults_for_non_docx(self, tmp_path: Path) -> None:
        txt_file = tmp_path / "not_docx.txt"
        txt_file.write_text("hello")
        info = _get_docx_page_info(str(txt_file))
        assert info["lines_per_page"] > 0
        assert info["chars_per_line"] > 0

    def test_parses_us_letter_dimensions(self, docx_file: Path) -> None:
        """python-docx creates US Letter (12240x15840) with 1-inch margins."""
        info = _get_docx_page_info(str(docx_file))
        # US Letter: 15840 - 1440*2 = 12960 usable height
        # 12960 / 253 ≈ 51 lines per page
        assert info["lines_per_page"] >= 40
        assert info["chars_per_line"] >= 60


# ─── _has_page_break ───────────────────────────────────────────────────────


class TestHasPageBreak:
    """Tests for ``_has_page_break``."""

    def test_returns_false_for_plain_element(self) -> None:
        """A simple paragraph without page breaks should return False."""
        try:
            from mammoth import documents
        except ImportError:
            pytest.skip("mammoth not installed")

        para = documents.paragraph(
            style_id="Normal",
            style_name="Normal",
            numbering=None,
            alignment=None,
            indent=None,
            children=[documents.run(children=[documents.text("hello")])],
        )
        assert _has_page_break(para) is False

    def test_detects_page_break_element(self) -> None:
        """A Break with break_type='page' should be detected."""
        try:
            from mammoth import documents
        except ImportError:
            pytest.skip("mammoth not installed")

        page_break = documents.page_break
        para = documents.paragraph(
            style_id="Normal",
            style_name="Normal",
            numbering=None,
            alignment=None,
            indent=None,
            children=[page_break],
        )
        assert _has_page_break(para) is True

    def test_detects_nested_page_break(self) -> None:
        """A page break inside a run inside a paragraph should be detected."""
        try:
            from mammoth import documents
        except ImportError:
            pytest.skip("mammoth not installed")

        run_with_break = documents.run(
            children=[documents.text("before"), documents.page_break]
        )
        para = documents.paragraph(
            style_id="Normal",
            style_name="Normal",
            numbering=None,
            alignment=None,
            indent=None,
            children=[run_with_break],
        )
        assert _has_page_break(para) is True


# ─── _estimate_element_lines ───────────────────────────────────────────────


class TestEstimateElementLines:
    """Tests for ``_estimate_element_lines``."""

    def test_empty_paragraph_returns_one(self) -> None:
        try:
            from mammoth import documents
        except ImportError:
            pytest.skip("mammoth not installed")

        para = documents.paragraph(
            style_id="Normal",
            style_name="Normal",
            numbering=None,
            alignment=None,
            indent=None,
            children=[],
        )
        assert _estimate_element_lines(para, 78) == 1.0

    def test_short_paragraph_returns_at_least_one(self) -> None:
        try:
            from mammoth import documents
        except ImportError:
            pytest.skip("mammoth not installed")

        para = documents.paragraph(
            style_id="Normal",
            style_name="Normal",
            numbering=None,
            alignment=None,
            indent=None,
            children=[documents.run(children=[documents.text("short")])],
        )
        result = _estimate_element_lines(para, 78)
        assert result >= 1.0

    def test_long_paragraph_spans_multiple_lines(self) -> None:
        try:
            from mammoth import documents
        except ImportError:
            pytest.skip("mammoth not installed")

        long_text = "x" * 500  # 500 chars / 78 chars_per_line ≈ 6.4 lines
        para = documents.paragraph(
            style_id="Normal",
            style_name="Normal",
            numbering=None,
            alignment=None,
            indent=None,
            children=[documents.run(children=[documents.text(long_text)])],
        )
        result = _estimate_element_lines(para, 78)
        assert result >= 5.0

    def test_heading_has_extra_lines(self) -> None:
        try:
            from mammoth import documents
        except ImportError:
            pytest.skip("mammoth not installed")

        text = "x" * 100  # ~1.3 text lines + 2 heading extra = ~3.3
        para = documents.paragraph(
            style_id="Heading1",
            style_name="Heading 1",
            numbering=None,
            alignment=None,
            indent=None,
            children=[documents.run(children=[documents.text(text)])],
        )
        result = _estimate_element_lines(para, 78)
        # Should be more than just text lines (100/78 ≈ 1.3)
        assert result > 1.3


# ─── _build_docx_transform ─────────────────────────────────────────────────


class TestBuildDocxTransform:
    """Tests for ``_build_docx_transform``."""

    def test_returns_callable(self) -> None:
        transform = _build_docx_transform(
            {"lines_per_page": 45, "chars_per_line": 78}
        )
        assert callable(transform)

    def test_inserts_marker_for_explicit_break(self) -> None:
        """Transform should insert a marker paragraph before an element with
        an explicit page break."""
        try:
            from mammoth import documents
        except ImportError:
            pytest.skip("mammoth not installed")

        transform = _build_docx_transform(
            {"lines_per_page": 45, "chars_per_line": 78}
        )

        para_with_break = documents.paragraph(
            style_id="Normal",
            style_name="Normal",
            numbering=None,
            alignment=None,
            indent=None,
            children=[documents.page_break],
        )
        doc = documents.document(
            notes=documents.notes(notes_list={}),
            comments=[],
            children=[para_with_break],
        )

        result = transform(doc)
        # Should have 2 children: marker paragraph + original
        assert len(result.children) == 2
        # First child should contain the marker text
        first = result.children[0]
        assert _has_page_break(first) is False
        # Check marker text is present
        marker_found = False

        def check_text(el):
            nonlocal marker_found
            if isinstance(el, documents.Text) and _PAGE_BREAK_MARKER in el.value:
                marker_found = True
            elif hasattr(el, "children"):
                for c in el.children:
                    check_text(c)

        check_text(first)
        assert marker_found

    def test_inserts_marker_for_estimated_overflow(self) -> None:
        """Transform should insert a marker when accumulated lines exceed
        lines_per_page."""
        try:
            from mammoth import documents
        except ImportError:
            pytest.skip("mammoth not installed")

        # Small page: 3 lines per page
        transform = _build_docx_transform(
            {"lines_per_page": 3, "chars_per_line": 78}
        )

        # Create 5 paragraphs, each ~1 line
        paras = []
        for i in range(5):
            paras.append(
                documents.paragraph(
                    style_id="Normal",
                    style_name="Normal",
                    numbering=None,
                    alignment=None,
                    indent=None,
                    children=[
                        documents.run(children=[documents.text(f"para {i}")])
                    ],
                )
            )

        doc = documents.document(
            notes=documents.notes(notes_list={}),
            comments=[],
            children=paras,
        )

        result = transform(doc)
        # After 3 paragraphs, the 4th should trigger a page break
        # So we expect at least 1 marker inserted
        markers = [
            c for c in result.children if _contains_marker(c)
        ]
        assert len(markers) >= 1

    def test_no_marker_for_short_document(self) -> None:
        """Transform should not insert markers for a short document."""
        try:
            from mammoth import documents
        except ImportError:
            pytest.skip("mammoth not installed")

        transform = _build_docx_transform(
            {"lines_per_page": 45, "chars_per_line": 78}
        )

        para = documents.paragraph(
            style_id="Normal",
            style_name="Normal",
            numbering=None,
            alignment=None,
            indent=None,
            children=[documents.run(children=[documents.text("short")])],
        )
        doc = documents.document(
            notes=documents.notes(notes_list={}),
            comments=[],
            children=[para],
        )

        result = transform(doc)
        assert len(result.children) == 1  # No marker added


def _contains_marker(element) -> bool:
    """Check if an element tree contains the page-break marker text."""
    try:
        from mammoth import documents
    except ImportError:
        return False

    def check(el):
        if isinstance(el, documents.Text) and _PAGE_BREAK_MARKER in el.value:
            return True
        if hasattr(el, "children"):
            return any(check(c) for c in el.children)
        return False

    return check(element)


# ─── _convert_docx_to_html ─────────────────────────────────────────────────


class TestConvertDocxToHtml:
    """Tests for ``_convert_docx_to_html``."""

    def test_converts_docx_to_html(self, docx_file: Path) -> None:
        html = _convert_docx_to_html(str(docx_file))
        assert isinstance(html, str)
        assert len(html) > 0
        assert "<p>" in html or "<h1>" in html

    def test_html_contains_content(self, docx_file: Path) -> None:
        html = _convert_docx_to_html(str(docx_file))
        assert "Hello World" in html
        assert "Test Document" in html or "Chapter 1" in html

    def test_long_docx_has_page_break_markers(
        self, long_docx_file: Path
    ) -> None:
        """A long document should have at least one page-break div."""
        html = _convert_docx_to_html(str(long_docx_file))
        assert 'class="docx-page-break"' in html

    def test_explicit_page_break_detected(
        self, explicit_break_docx: Path
    ) -> None:
        """A document with an explicit page break should produce a
        page-break div."""
        html = _convert_docx_to_html(str(explicit_break_docx))
        assert 'class="docx-page-break"' in html

    def test_no_marker_in_short_doc(self, docx_file: Path) -> None:
        """A short document should not have page-break markers."""
        html = _convert_docx_to_html(str(docx_file))
        assert 'class="docx-page-break"' not in html

    def test_unsupported_extension_raises_415(self, tmp_path: Path) -> None:
        from fastapi import HTTPException

        txt_file = tmp_path / "test.txt"
        txt_file.write_text("hello")
        with pytest.raises(HTTPException) as exc_info:
            _convert_docx_to_html(str(txt_file))
        assert exc_info.value.status_code == 415

    def test_old_doc_format_raises_415(self, tmp_path: Path) -> None:
        from fastapi import HTTPException

        doc_file = tmp_path / "old.doc"
        doc_file.write_bytes(b"fake doc content")
        with pytest.raises(HTTPException) as exc_info:
            _convert_docx_to_html(str(doc_file))
        assert exc_info.value.status_code == 415

    def test_no_residual_marker_in_output(self, docx_file: Path) -> None:
        """The marker sentinel should never appear in the final HTML."""
        html = _convert_docx_to_html(str(docx_file))
        assert _PAGE_BREAK_MARKER not in html
