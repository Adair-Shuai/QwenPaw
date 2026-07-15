# -*- coding: utf-8 -*-
"""Integration tests for the workspace /convert-office endpoint.

Tests cover:
- DOCX → HTML conversion via HTTP
- Page-break marker insertion in the API response
- Error handling for missing files and unsupported types
- Agent-scoped file resolution (X-Agent-Id header)
"""
from __future__ import annotations

import pytest


@pytest.mark.integration
@pytest.mark.p1
def test_convert_office_docx_returns_html(app_server) -> None:
    """Test purpose:
    - Verify the /api/workspace/convert-office endpoint converts a .docx file
      to HTML.

    Test flow:
    1. Create a test agent.
    2. Write a .docx file into the agent workspace using python-docx.
    3. POST /api/workspace/convert-office with the file URL.
    4. Assert response is 200 and contains HTML.
    5. Delete test agent.

    API endpoints:
    - POST /api/agents
    - POST /api/workspace/convert-office
    - DELETE /api/agents/{agentId}
    """
    agent_id = "integ_convert_01"
    headers = {"X-Agent-Id": agent_id}

    create_agent = app_server.api_request(
        "POST",
        "/api/agents",
        json={
            "id": agent_id,
            "name": "Convert office test agent",
            "description": "",
        },
    )
    assert create_agent.status_code == 201, app_server.logs_tail()

    try:
        # Create a .docx file in the workspace
        from docx import Document

        docx_path = (
            app_server.working_dir
            / "workspaces"
            / agent_id
            / "test_convert.docx"
        )
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_heading("Integration Test Document", level=0)
        doc.add_paragraph("This is a test paragraph for convert-office.")
        doc.save(str(docx_path))

        # Call the convert-office endpoint
        resp = app_server.api_request(
            "POST",
            "/api/workspace/convert-office",
            headers=headers,
            json={
                "url": "/api/workspace/binary-files/test_convert.docx",
                "mimeType": "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document",
            },
        )
        assert resp.status_code == 200, app_server.logs_tail()
        data = resp.json()
        assert "html" in data
        assert isinstance(data["html"], str)
        assert len(data["html"]) > 0
        assert "Integration Test Document" in data["html"]
        assert "test paragraph" in data["html"]
    finally:
        app_server.api_request("DELETE", f"/api/agents/{agent_id}")


@pytest.mark.integration
@pytest.mark.p1
def test_convert_office_long_docx_has_page_breaks(app_server) -> None:
    """Test purpose:
    - Verify the convert-office endpoint inserts page-break markers for a
      long document.

    Test flow:
    1. Create a test agent.
    2. Write a long .docx file (>1 page of content) into the workspace.
    3. POST /api/workspace/convert-office.
    4. Assert response HTML contains page-break divs.
    5. Delete test agent.

    API endpoints:
    - POST /api/agents
    - POST /api/workspace/convert-office
    - DELETE /api/agents/{agentId}
    """
    agent_id = "integ_convert_pb_01"
    headers = {"X-Agent-Id": agent_id}

    create_agent = app_server.api_request(
        "POST",
        "/api/agents",
        json={
            "id": agent_id,
            "name": "Convert office page break agent",
            "description": "",
        },
    )
    assert create_agent.status_code == 201, app_server.logs_tail()

    try:
        from docx import Document

        docx_path = (
            app_server.working_dir / "workspaces" / agent_id / "long_doc.docx"
        )
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_heading("Long Document", level=0)
        for i in range(60):
            doc.add_paragraph(f"Paragraph {i} " * 10)
        doc.save(str(docx_path))

        resp = app_server.api_request(
            "POST",
            "/api/workspace/convert-office",
            headers=headers,
            json={
                "url": "/api/workspace/binary-files/long_doc.docx",
            },
        )
        assert resp.status_code == 200, app_server.logs_tail()
        html = resp.json()["html"]
        assert 'class="docx-page-break"' in html
    finally:
        app_server.api_request("DELETE", f"/api/agents/{agent_id}")


@pytest.mark.integration
@pytest.mark.p2
def test_convert_office_missing_file_returns_404(app_server) -> None:
    """Test purpose:
    - Verify convert-office returns 404 for a nonexistent file.

    Test flow:
    1. Create a test agent.
    2. POST /api/workspace/convert-office with a non-existent file URL.
    3. Assert 404 status.
    4. Delete test agent.

    API endpoints:
    - POST /api/agents
    - POST /api/workspace/convert-office
    - DELETE /api/agents/{agentId}
    """
    agent_id = "integ_convert_404_01"
    headers = {"X-Agent-Id": agent_id}

    create_agent = app_server.api_request(
        "POST",
        "/api/agents",
        json={
            "id": agent_id,
            "name": "Convert 404 agent",
            "description": "",
        },
    )
    assert create_agent.status_code == 201, app_server.logs_tail()

    try:
        resp = app_server.api_request(
            "POST",
            "/api/workspace/convert-office",
            headers=headers,
            json={"url": "/api/workspace/binary-files/nonexistent.docx"},
        )
        assert resp.status_code == 404, app_server.logs_tail()
    finally:
        app_server.api_request("DELETE", f"/api/agents/{agent_id}")


@pytest.mark.integration
@pytest.mark.p2
def test_convert_office_unsupported_type_returns_415(app_server) -> None:
    """Test purpose:
    - Verify convert-office returns 415 for unsupported file types.

    Test flow:
    1. Create a test agent.
    2. Write a .txt file into the workspace.
    3. POST /api/workspace/convert-office with the .txt file URL.
    4. Assert 415 status.
    5. Delete test agent.

    API endpoints:
    - POST /api/agents
    - POST /api/workspace/convert-office
    - DELETE /api/agents/{agentId}
    """
    agent_id = "integ_convert_415_01"
    headers = {"X-Agent-Id": agent_id}

    create_agent = app_server.api_request(
        "POST",
        "/api/agents",
        json={
            "id": agent_id,
            "name": "Convert 415 agent",
            "description": "",
        },
    )
    assert create_agent.status_code == 201, app_server.logs_tail()

    try:
        # Write a .txt file
        txt_path = (
            app_server.working_dir / "workspaces" / agent_id / "plain.txt"
        )
        txt_path.parent.mkdir(parents=True, exist_ok=True)
        txt_path.write_text("plain text content")

        resp = app_server.api_request(
            "POST",
            "/api/workspace/convert-office",
            headers=headers,
            json={"url": "/api/workspace/binary-files/plain.txt"},
        )
        assert resp.status_code == 415, app_server.logs_tail()
    finally:
        app_server.api_request("DELETE", f"/api/agents/{agent_id}")
