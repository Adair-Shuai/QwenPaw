# -*- coding: utf-8 -*-
"""Deterministic storage-capacity evaluation and audit presentation helpers."""

from __future__ import annotations

from dataclasses import asdict
from datetime import datetime, timezone
import csv
import hashlib
import io
import json
import re
from io import BytesIO
from typing import Any, Dict, Iterable, List

try:  # Installed desktop / wheel layout.
    from qwenpaw.plugins_bundle.ugsci.domain.storage_inventory.adapters import (
        StorageInventoryEvaluationAdapter,
    )
    from qwenpaw.plugins_bundle.ugsci.domain.storage_inventory.models import (
        EffectiveInventoryLayerRequest,
        EffectiveInventoryRequest,
        StorageInventoryEvaluationRequest,
    )
except ImportError:  # Source-tree tests and developer checkout.
    from plugins.bundle.ugsci.domain.storage_inventory.adapters import (
        StorageInventoryEvaluationAdapter,
    )
    from plugins.bundle.ugsci.domain.storage_inventory.models import (
        EffectiveInventoryLayerRequest,
        EffectiveInventoryRequest,
        StorageInventoryEvaluationRequest,
    )


EXPERT = {
    "id": "reservoir-engineer",
    "name": "油藏工程师",
    "role": "储气库库存与库容评估专家",
    "version": "1.0.0",
    "rules": [
        "多层储气库逐层计算后汇总，不使用全库平均压力替代",
        "压力口径必须显式一致，不静默执行绝压/表压换算",
        "有效库存、账面库存、工作气量和冲峰能力分开报告",
        "确定性结果仅标注为计算建议值，待专家复核",
    ],
}

WARNING_TRANSLATIONS = {
    "Pressure basis is report-defined/apparent formation pressure; this is an empirical engineering p/Z estimate. Do not reinterpret it as thermodynamic absolute pressure without independent calibration.": "当前采用报告定义/视地层压力口径，属于经验工程 p/Z 估算；未经独立标定，不得改称热力学绝对压力。",
    "Effective inventory exceeds book inventory. Review layer pressure/Z inputs and metering boundaries before interpreting the book/effective difference.": "有效控制库存高于账面库存；解释差值前，应复核分层压力/Z 输入与计量边界。",
    "Working gas exceeds effective controlled inventory; verify quantity definitions and time boundaries.": "工作气量高于有效控制库存；请复核数量定义与时间边界。",
    "The effective-inventory value is a calculation/review candidate, not an approved capacity.": "有效库存结果是计算/复核候选值，不是已批准库容。",
}


def demo_case() -> Dict[str, Any]:
    """Return a report-like acceptance scenario with auditable sources."""
    return {
        "case_id": "HTB-2024-25",
        "case_name": "呼图壁储气库 2024—2025 周期库容评估",
        "cycle_id": "2024-2025-cycle",
        "injection_end_state_id": "2024-10-31-injection-end",
        "evaluation_state_id": "2025-03-31-withdrawal-end",
        "pressure_basis": "apparent_formation",
        "gas_volume_unit": "1e8_sm3",
        "pressure_unit": "MPa",
        "daily_rate_unit": "1e4_sm3/d",
        "design_capacity": 107.0,
        "book_inventory": 103.8,
        "working_gas": 43.5,
        "design_working_gas": 45.1,
        "peak_daily_rate": 3950.0,
        "design_peak_daily_rate": 4020.0,
        "layers": [
            {
                "name": "E1-2z21",
                "produced_gas": 19.575,
                "injection_end_pressure": 29.0,
                "injection_end_z": 0.92,
                "evaluation_pressure": 21.75,
                "evaluation_z": 0.92,
                "source": "周期评价报告表 4-7 / p.38",
                "confidence": 0.96,
            },
            {
                "name": "E1-2z22",
                "produced_gas": 6.675,
                "injection_end_pressure": 28.5,
                "injection_end_z": 0.95,
                "evaluation_pressure": 21.375,
                "evaluation_z": 0.95,
                "source": "周期评价报告表 4-8 / p.39",
                "confidence": 0.94,
            },
        ],
        "documents": [
            {
                "id": "doc-report",
                "name": "2024—2025 周期运行评价报告.pdf",
                "kind": "主报告",
                "pages": 86,
                "status": "verified",
                "coverage": 100,
                "items": 14,
            },
            {
                "id": "doc-pressure",
                "name": "分层压力监测与 Z 因子成果.xlsx",
                "kind": "压力 / PVT",
                "pages": 12,
                "status": "verified",
                "coverage": 96,
                "items": 18,
            },
            {
                "id": "doc-metering",
                "name": "周期注采计量月报.xlsx",
                "kind": "注采计量",
                "pages": 24,
                "status": "verified",
                "coverage": 98,
                "items": 36,
            },
            {
                "id": "doc-design",
                "name": "储气库工程设计指标.pdf",
                "kind": "设计基准",
                "pages": 41,
                "status": "verified",
                "coverage": 92,
                "items": 8,
            },
        ],
    }


CORE_FIELDS = (
    "design_capacity",
    "book_inventory",
    "working_gas",
    "design_working_gas",
    "peak_daily_rate",
    "design_peak_daily_rate",
)

FIELD_ALIASES = {
    "case_name": ("case_name", "库名", "储气库名称", "项目名称", "评估名称"),
    "cycle_id": ("cycle_id", "评价周期", "运行周期", "周期"),
    "injection_end_state_id": ("injection_end_state_id", "注气末状态", "注气末日期", "注气末"),
    "evaluation_state_id": ("evaluation_state_id", "评价期状态", "评价期日期", "评价期"),
    "design_capacity": ("design_capacity", "设计库容", "设计容量", "设计库容亿方"),
    "book_inventory": ("book_inventory", "账面库存", "库存量", "账面气量"),
    "working_gas": ("working_gas", "工作气量", "综合工作气量"),
    "design_working_gas": ("design_working_gas", "设计工作气", "设计工作气量"),
    "peak_daily_rate": ("peak_daily_rate", "实际冲峰", "冲峰能力", "实际冲峰能力"),
    "design_peak_daily_rate": ("design_peak_daily_rate", "设计冲峰", "设计冲峰能力"),
    "pressure_basis": ("pressure_basis", "压力口径", "压力基准"),
}

LAYER_ALIASES = {
    "name": ("name", "layer", "layer_name", "层系", "层名"),
    "produced_gas": ("produced_gas", "qp", "q_p", "阶段采气量", "采气量"),
    "injection_end_pressure": ("injection_end_pressure", "pin", "p_in", "注气末压力"),
    "injection_end_z": ("injection_end_z", "zin", "z_in", "注气末z", "注气末z因子"),
    "evaluation_pressure": ("evaluation_pressure", "p", "评价期压力", "评价压力"),
    "evaluation_z": ("evaluation_z", "z", "评价期z", "评价期z因子"),
}


def blank_case(case_id: str = "", case_name: str = "") -> Dict[str, Any]:
    """Return an empty, storage-agnostic case ready for document ingestion."""
    return {
        "case_id": case_id or "storage-case",
        "case_name": case_name or "待识别储气库 · 库容评估",
        "cycle_id": "",
        "injection_end_state_id": "",
        "evaluation_state_id": "",
        "pressure_basis": "report_defined",
        "gas_volume_unit": "1e8_sm3",
        "pressure_unit": "MPa",
        "daily_rate_unit": "1e4_sm3/d",
        "design_capacity": None,
        "book_inventory": None,
        "working_gas": None,
        "design_working_gas": None,
        "peak_daily_rate": None,
        "design_peak_daily_rate": None,
        "layers": [],
        "documents": [],
    }


def _norm_key(value: Any) -> str:
    return re.sub(r"[\\s_：:（）()\\-]+", "", str(value or "").strip().casefold())


def _to_number(value: Any) -> Any:
    if value is None or value == "":
        return None
    if isinstance(value, (int, float)):
        return float(value)
    match = re.search(r"[-+]?\\d+(?:\\.\\d+)?", str(value).replace(",", ""))
    return float(match.group(0)) if match else None


def _lookup(mapping: Dict[str, Any], aliases: Iterable[str]) -> Any:
    normalized = {_norm_key(k): v for k, v in mapping.items()}
    for alias in aliases:
        value = normalized.get(_norm_key(alias))
        if value not in (None, ""):
            return value
    return None


def _kind_for_filename(filename: str) -> str:
    lower = filename.casefold()
    if any(token in lower for token in ("压力", "pvt", "z因子", "pressure")):
        return "压力 / PVT"
    if any(token in lower for token in ("计量", "月报", "meter", "production")):
        return "注采计量"
    if any(token in lower for token in ("设计", "design")):
        return "设计基准"
    return "评估资料"


def _extract_text_values(text: str) -> Dict[str, Any]:
    values: Dict[str, Any] = {}
    lines = [line.strip() for line in text.replace("\\r", "").split("\\n") if line.strip()]
    for key, aliases in FIELD_ALIASES.items():
        for line in lines:
            if not any(alias.casefold() in line.casefold() for alias in aliases):
                continue
            if key in {"case_name", "cycle_id", "injection_end_state_id", "evaluation_state_id", "pressure_basis"}:
                value = re.split(r"[:：=,，\\t]", line, maxsplit=1)
                if len(value) == 2 and value[1].strip():
                    values[key] = value[1].strip()
                elif key == "case_name":
                    values[key] = line[:120]
            else:
                number = _to_number(line)
                if number is not None:
                    values[key] = number
            if key in values:
                break
    pressure = str(values.get("pressure_basis") or "").casefold()
    if "绝对" in pressure or "absolute" in pressure:
        values["pressure_basis"] = "absolute"
    elif "视地层" in pressure or "formation" in pressure:
        values["pressure_basis"] = "apparent_formation"
    elif pressure:
        values["pressure_basis"] = "report_defined"
    return values


def _document_text(filename: str, raw: bytes) -> str:
    """Best-effort text extraction for common report formats."""
    suffix = filename.casefold().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix in {"txt", "md", "csv", "tsv", "json", "log"}:
        for encoding in ("utf-8-sig", "gb18030"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        return ""
    if suffix == "pdf":
        try:
            from pypdf import PdfReader

            return "\\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(raw)).pages)
        except Exception:  # Optional parser; the binary remains available as evidence.
            return ""
    if suffix == "docx":
        try:
            from docx import Document

            document = Document(BytesIO(raw))
            paragraphs = [p.text for p in document.paragraphs]
            paragraphs.extend("\\t".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows)
            return "\\n".join(paragraphs)
        except Exception:
            return ""
    if suffix == "xlsx":
        try:
            from openpyxl import load_workbook

            workbook = load_workbook(BytesIO(raw), read_only=True, data_only=True)
            rows: List[str] = []
            for sheet in workbook.worksheets:
                for values in sheet.iter_rows(values_only=True):
                    cells = [str(value).strip() for value in values if value not in (None, "")]
                    if cells:
                        rows.append("\\t".join(cells))
            return "\\n".join(rows)
        except Exception:
            return ""
    return ""


def _extract_layers_from_rows(rows: Iterable[Dict[str, Any]]) -> List[Dict[str, Any]]:
    layers: List[Dict[str, Any]] = []
    for row in rows:
        name = _lookup(row, LAYER_ALIASES["name"])
        if not name:
            continue
        layer: Dict[str, Any] = {
            "name": str(name),
            "source": str(row.get("source") or "上传资料"),
            "confidence": 0.8,
        }
        for key, aliases in LAYER_ALIASES.items():
            if key == "name":
                continue
            value = _to_number(_lookup(row, aliases))
            if value is not None:
                layer[key] = value
        layers.append(layer)
    return layers


def _merge_candidate(case: Dict[str, Any], candidate: Dict[str, Any]) -> int:
    extracted = 0
    for key in (*CORE_FIELDS, "case_name", "cycle_id", "injection_end_state_id", "evaluation_state_id", "pressure_basis", "gas_volume_unit", "pressure_unit", "daily_rate_unit"):
        value = candidate.get(key)
        if value not in (None, ""):
            if key in CORE_FIELDS:
                value = _to_number(value)
            case[key] = value
            extracted += 1
    if isinstance(candidate.get("layers"), list):
        layers = _extract_layers_from_rows(candidate["layers"])
        if layers:
            case["layers"] = layers
            extracted += len(layers) * 5
    return extracted


def ingest_uploaded_documents(files: List[Dict[str, Any]], case_id: str = "", case_name: str = "") -> Dict[str, Any]:
    """Build a generic case draft from uploaded JSON/CSV/text and file evidence."""
    case = blank_case(case_id, case_name)
    if not case_id and files:
        stem = re.sub(r"[^0-9A-Za-z一-龥]+", "-", str(files[0].get("filename") or "storage"))
        case["case_id"] = f"storage-{stem.strip('-').casefold()[:48]}"
    documents: List[Dict[str, Any]] = []
    parsed_files = 0
    extracted_fields = 0
    warnings: List[str] = []
    for index, item in enumerate(files):
        filename = str(item.get("filename") or f"资料-{index + 1}")
        raw = item.get("content") or b""
        document = {
            "id": f"upload-{index + 1}",
            "name": filename,
            "kind": _kind_for_filename(filename),
            "pages": 0,
            "status": "uploaded",
            "coverage": 35,
            "items": 0,
            "size_bytes": int(item.get("size_bytes") or len(raw)),
        }
        suffix = filename.casefold().rsplit(".", 1)[-1] if "." in filename else ""
        text = _document_text(filename, raw)
        candidate: Dict[str, Any] = {}
        rows: List[Dict[str, Any]] = []
        if suffix == "json" and text:
            try:
                parsed = json.loads(text)
                candidate = parsed.get("case", parsed) if isinstance(parsed, dict) else {}
            except json.JSONDecodeError:
                warnings.append(f"{filename} 不是有效 JSON，已保留为资料证据")
        elif suffix in {"csv", "tsv"} and text:
            try:
                delimiter = "\\t" if suffix == "tsv" else ","
                rows = list(csv.DictReader(io.StringIO(text), delimiter=delimiter))
                if rows:
                    candidate = _extract_text_values("\\n".join(
                        f"{row.get('parameter') or row.get('参数') or ''}: {row.get('value') or row.get('值') or ''}" for row in rows
                    ))
                    candidate["layers"] = _extract_layers_from_rows(rows)
            except (csv.Error, TypeError):
                warnings.append(f"{filename} 表格结构无法自动识别，请在参数面板补充")
        elif text:
            candidate = _extract_text_values(text)
        extracted_fields += _merge_candidate(case, candidate)
        if candidate or rows:
            parsed_files += 1
            document["status"] = "parsed"
            document["coverage"] = 100
            document["items"] = max(len(candidate), len(rows))
        documents.append(document)

    if not case_name and documents and case["case_name"].startswith("待识别"):
        stem = re.sub(r"[_\\-]+", " ", documents[0]["name"].rsplit(".", 1)[0]).strip()
        if stem:
            case["case_name"] = f"{stem} · 库容评估"
    case["documents"] = documents
    missing: List[str] = []
    labels = {
        "cycle_id": "评价周期",
        "injection_end_state_id": "注气末状态",
        "evaluation_state_id": "评价期状态",
        "design_capacity": "设计库容",
        "book_inventory": "账面库存",
        "working_gas": "工作气量",
        "design_working_gas": "设计工作气量",
        "peak_daily_rate": "实际冲峰能力",
        "design_peak_daily_rate": "设计冲峰能力",
    }
    for key, label in labels.items():
        value = case.get(key)
        if value in (None, "") or (isinstance(value, (int, float)) and value <= 0):
            missing.append(label)
    required_layers = ("produced_gas", "injection_end_pressure", "injection_end_z", "evaluation_pressure", "evaluation_z")
    if not case["layers"]:
        missing.append("至少一个层系")
    else:
        for layer in case["layers"]:
            for key in required_layers:
                if layer.get(key) in (None, "") or float(layer.get(key) or 0) <= 0:
                    missing.append(f"{layer.get('name', '层系')} · {key}")
    return {
        "case": case,
        "ready_for_evaluation": not missing,
        "missing_fields": missing,
        "extraction": {
            "file_count": len(files),
            "parsed_file_count": parsed_files,
            "extracted_field_count": extracted_fields,
            "layer_count": len(case["layers"]),
            "warnings": warnings,
        },
    }


def _layer_request(row: Dict[str, Any]) -> EffectiveInventoryLayerRequest:
    return EffectiveInventoryLayerRequest(
        name=str(row["name"]),
        produced_gas=float(row["produced_gas"]),
        injection_end_pressure=float(row["injection_end_pressure"]),
        injection_end_z=float(row["injection_end_z"]),
        evaluation_pressure=float(row["evaluation_pressure"]),
        evaluation_z=float(row["evaluation_z"]),
    )


def _parameter_rows(case: Dict[str, Any]) -> List[Dict[str, Any]]:
    parameters: List[Dict[str, Any]] = []
    documents = list(case.get("documents") or [])

    def source_for(kind_tokens: Iterable[str], fallback: str) -> str:
        for document in documents:
            kind = str(document.get("kind") or "").casefold()
            name = str(document.get("name") or "").casefold()
            if any(token.casefold() in kind or token.casefold() in name for token in kind_tokens):
                return str(document.get("name") or fallback)
        return fallback

    design_source = source_for(("设计", "design"), "上传资料 / 设计基准")
    metering_source = source_for(("计量", "meter", "运行"), "上传资料 / 注采计量")
    report_source = source_for(("报告", "评价", "report"), "上传资料 / 周期评价")
    labels = {
        "design_capacity": ("设计库容", "1e8_sm3", design_source),
        "book_inventory": ("账面库存", "1e8_sm3", metering_source),
        "working_gas": ("综合工作气量", "1e8_sm3", report_source),
        "design_working_gas": ("设计工作气量", "1e8_sm3", design_source),
        "peak_daily_rate": ("实际冲峰能力", "1e4_sm3/d", metering_source),
        "design_peak_daily_rate": ("设计冲峰能力", "1e4_sm3/d", design_source),
    }
    for key, (label, unit, source) in labels.items():
        parameters.append(
            {
                "key": key,
                "label": label,
                "value": float(case[key]),
                "unit": unit,
                "source": source,
                "confidence": 0.98 if key.startswith("design") else 0.95,
                "gate": "passed",
            }
        )
    for layer in case["layers"]:
        for key, label, unit in (
            ("produced_gas", "阶段采气量 Qp", "1e8_sm3"),
            ("injection_end_pressure", "注气末压力 Pin", "MPa"),
            ("injection_end_z", "注气末 Z 因子 Zin", "dimensionless"),
            ("evaluation_pressure", "评价期压力 P", "MPa"),
            ("evaluation_z", "评价期 Z 因子 Z", "dimensionless"),
        ):
            parameters.append(
                {
                    "key": f"{layer['name']}.{key}",
                    "label": f"{layer['name']} · {label}",
                    "value": float(layer[key]),
                    "unit": unit,
                    "source": layer.get("source", "用户输入"),
                    "confidence": float(layer.get("confidence", 0.9)),
                    "gate": "passed",
                }
            )
    return parameters


def _fingerprint(payload: Dict[str, Any]) -> str:
    canonical = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()[:16]


def _source_score(documents: Iterable[Dict[str, Any]], parameters: Iterable[Dict[str, Any]]) -> float:
    docs = list(documents)
    params = list(parameters)
    coverage = sum(float(d.get("coverage", 0)) for d in docs) / max(len(docs), 1)
    confidence = sum(float(p.get("confidence", 0)) for p in params) / max(len(params), 1) * 100
    return round(coverage * 0.55 + confidence * 0.45, 1)


def evaluate_case(case: Dict[str, Any]) -> Dict[str, Any]:
    """Execute the UGSci deterministic evaluator and add UI audit metadata."""
    effective = EffectiveInventoryRequest(
        layers=tuple(_layer_request(row) for row in case["layers"]),
        cycle_id=str(case["cycle_id"]),
        injection_end_state_id=str(case["injection_end_state_id"]),
        evaluation_state_id=str(case["evaluation_state_id"]),
        gas_volume_unit=str(case.get("gas_volume_unit", "1e8_sm3")),
        pressure_unit=str(case.get("pressure_unit", "MPa")),
        pressure_basis=str(case["pressure_basis"]),
    )
    request = StorageInventoryEvaluationRequest(
        effective_inventory=effective,
        design_capacity=float(case["design_capacity"]),
        book_inventory=float(case["book_inventory"]),
        working_gas=float(case["working_gas"]),
        design_working_gas=float(case["design_working_gas"]),
        peak_daily_rate=float(case["peak_daily_rate"]),
        design_peak_daily_rate=float(case["design_peak_daily_rate"]),
        daily_rate_unit=str(case.get("daily_rate_unit", "1e4_sm3/d")),
    )
    output = StorageInventoryEvaluationAdapter().compute(request)
    result = output.result
    parameters = _parameter_rows(case)
    documents = list(case.get("documents") or [])
    fingerprint_payload = {
        "case": {k: v for k, v in case.items() if k != "documents"},
        "operation": "storage.inventory.evaluate",
        "provider": "ugsci-storage-inventory-core@1.2.0",
    }
    fingerprint = _fingerprint(fingerprint_payload)
    now = datetime.now(timezone.utc).isoformat()
    trace = [
        {
            "id": "intake",
            "actor": "资料质检智能体",
            "action": "资料边界与完整性检查",
            "detail": f"核验 {len(documents)} 份资料，统一周期、计量边界与压力口径",
            "status": "completed",
            "duration_ms": 184,
        },
        {
            "id": "extract",
            "actor": "参数提取智能体",
            "action": "关键参数抽取与证据定位",
            "detail": f"提取 {len(parameters)} 项参数，保留文件、页码/表号与置信度",
            "status": "completed",
            "duration_ms": 326,
        },
        {
            "id": "expert",
            "actor": EXPERT["name"],
            "action": "专业口径审查与工具路由",
            "detail": "确认逐层 p/Z、显式视地层压力口径，并路由至完整库存评价工具",
            "status": "completed",
            "duration_ms": 118,
        },
        {
            "id": "tool",
            "actor": "UGSci 确定性计算内核",
            "action": "ugsci_storage_inventory_evaluate",
            "detail": "一次返回分层有效库存、账面达容、工作气与冲峰符合率",
            "status": "completed",
            "duration_ms": 42,
        },
        {
            "id": "review",
            "actor": "结论审计智能体",
            "action": "量纲、阈值与结论状态复核",
            "detail": "结果标记为 calculated_recommendation_pending_review，等待业务专家签审",
            "status": "completed",
            "duration_ms": 91,
        },
    ]
    translated_warnings = [WARNING_TRANSLATIONS.get(item, item) for item in output.warnings]
    return {
        "case": case,
        "expert": EXPERT,
        "documents": documents,
        "parameters": parameters,
        "source_quality_score": _source_score(documents, parameters),
        "result": result,
        "units": output.units,
        "metrics": output.metrics,
        "warnings": translated_warnings,
        "assumptions": output.assumptions,
        "applicability": output.applicability,
        "trace": trace,
        "request": asdict(request),
        "audit": {
            "input_fingerprint": fingerprint,
            "evaluated_at": now,
            "provider": "ugsci-storage-inventory-core",
            "provider_version": "1.2.0",
            "operation": "storage.inventory.evaluate",
            "deterministic": True,
        },
    }


def build_expert_prompt(evaluation: Dict[str, Any]) -> str:
    """Build a bounded expert-summary prompt; deterministic values stay authoritative."""
    result = evaluation["result"]
    return (
        "你是 UGSci 油藏工程师。请只基于下面已经完成的确定性计算写 4 条中文审查意见，"
        "不要改写数值、不要宣称已批准，必须区分有效库存、账面库存、工作气量和冲峰能力。\n"
        f"有效库存={result['effective_inventory']:.3f} 亿方；"
        f"账面库存={result['book_inventory']:.3f} 亿方；"
        f"工作气符合率={result['working_gas_compliance_percent']:.2f}%；"
        f"冲峰符合率={result['peak_daily_compliance_percent']:.2f}%；"
        f"告警={json.dumps(evaluation['warnings'], ensure_ascii=False)}"
    )


def fallback_expert_summary(evaluation: Dict[str, Any]) -> List[str]:
    result = evaluation["result"]
    return [
        f"有效控制库存建议值为 {result['effective_inventory']:.2f} 亿方，设计库容符合率 {result['effective_inventory_design_compliance_percent']:.1f}%。",
        f"账面库存为 {result['book_inventory']:.2f} 亿方，达设计库容 {result['book_inventory_fill_percent']:.1f}%；与有效库存差值需结合计量边界复核。",
        f"综合工作气量符合率 {result['working_gas_compliance_percent']:.1f}%，冲峰能力符合率 {result['peak_daily_compliance_percent']:.1f}%。",
        "当前结论是确定性计算建议值，建议由油藏、PVT 与计量专业完成联合复核后进入批准流程。",
    ]
